import os
import re
from typing import List
import docx
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import tempfile

class ResumeEnhancer:
    def __init__(self):
        self.temp_dir = tempfile.mkdtemp()
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF file"""
        try:
            reader = PdfReader(pdf_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(docx_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")
    
    def extract_text_from_doc(self, doc_path: str) -> str:
        """Extract text from DOC file (treating as DOCX for simplicity)"""
        # Note: For .doc files, we'll try to read as docx
        # For production, you might want to use python-docx2txt or convert .doc to .docx
        try:
            return self.extract_text_from_docx(doc_path)
        except:
            raise Exception("DOC file format not fully supported. Please convert to DOCX or PDF.")
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from resume file based on extension"""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif ext == '.docx':
            return self.extract_text_from_docx(file_path)
        elif ext == '.doc':
            return self.extract_text_from_doc(file_path)
        else:
            raise Exception(f"Unsupported file format: {ext}")
    
    def extract_keywords(self, text: str) -> List[str]:
        """Extract important keywords from text"""
        # Convert to lowercase and split
        words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
        
        # Common stop words to filter out
        stop_words = {
            'the', 'and', 'for', 'are', 'but', 'not', 'you', 'all', 'can', 'her',
            'was', 'one', 'our', 'out', 'day', 'get', 'has', 'him', 'his', 'how',
            'its', 'may', 'new', 'now', 'old', 'see', 'two', 'way', 'who', 'boy',
            'did', 'she', 'use', 'her', 'many', 'than', 'them', 'these', 'with',
            'this', 'that', 'from', 'have', 'been', 'will', 'your', 'work', 'more'
        }
        
        # Filter and get unique keywords
        keywords = [w for w in words if w not in stop_words and len(w) > 3]
        
        # Count frequency
        from collections import Counter
        keyword_counts = Counter(keywords)
        
        # Return top keywords
        return [word for word, count in keyword_counts.most_common(30)]
    
    def enhance_resume_text(self, resume_text: str, job_posting: str) -> str:
        """Enhance resume text based on job posting keywords"""
        # Extract keywords from job posting
        job_keywords = self.extract_keywords(job_posting)
        resume_keywords = self.extract_keywords(resume_text)
        
        # Find missing keywords that should be added
        missing_keywords = [kw for kw in job_keywords if kw not in resume_keywords]
        
        # Enhance the resume by adding relevant keywords in context
        enhanced_text = resume_text
        
        # Try to intelligently add missing keywords
        # This is a simplified version - in production, you'd want more sophisticated NLP
        if missing_keywords:
            # Add a skills section if it doesn't exist and we have missing keywords
            if 'skills' not in resume_text.lower() and 'skill' not in resume_text.lower():
                skills_section = "\n\nSKILLS:\n" + ", ".join(missing_keywords[:10]) + "\n"
                enhanced_text += skills_section
            else:
                # Try to add to existing skills section
                enhanced_text = self.add_to_skills_section(enhanced_text, missing_keywords[:10])
        
        return enhanced_text
    
    def add_to_skills_section(self, text: str, keywords: List[str]) -> str:
        """Add keywords to existing skills section"""
        # Find skills section
        lines = text.split('\n')
        enhanced_lines = []
        skills_found = False
        skip_next = False
        
        for i, line in enumerate(lines):
            if skip_next:
                skip_next = False
                continue
                
            enhanced_lines.append(line)
            if ('skills' in line.lower() or 'skill' in line.lower()) and not skills_found:
                skills_found = True
                # Add keywords to next few lines
                if i + 1 < len(lines) and lines[i + 1].strip():
                    # Add to existing line
                    enhanced_lines.append(lines[i + 1] + ", " + ", ".join(keywords))
                    skip_next = True
                else:
                    # Add new line
                    enhanced_lines.append(", ".join(keywords))
        
        return '\n'.join(enhanced_lines)
    
    def create_pdf(self, text: str, output_path: str):
        """Create PDF from text"""
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Custom style for resume
        resume_style = ParagraphStyle(
            'ResumeStyle',
            parent=styles['Normal'],
            fontSize=11,
            leading=14,
            spaceAfter=6
        )
        
        # Split text into paragraphs
        paragraphs = text.split('\n')
        for para in paragraphs:
            if para.strip():
                p = Paragraph(para.replace('&', '&amp;'), resume_style)
                story.append(p)
                story.append(Spacer(1, 0.1*inch))
        
        doc.build(story)
    
    def create_docx(self, text: str, output_path: str):
        """Create DOCX from text"""
        doc = Document()
        
        # Split text into paragraphs
        paragraphs = text.split('\n')
        for para in paragraphs:
            if para.strip():
                # Check if it's a heading (all caps or short line)
                if para.isupper() and len(para) < 50:
                    p = doc.add_heading(para, level=2)
                else:
                    p = doc.add_paragraph(para)
                    p.style.font.size = Pt(11)
        
        doc.save(output_path)
    
    def enhance_resume(self, resume_path: str, job_posting: str, output_format: str = 'pdf') -> str:
        """Main method to enhance resume"""
        # Extract text from resume
        resume_text = self.extract_text(resume_path)
        
        # Enhance resume
        enhanced_text = self.enhance_resume_text(resume_text, job_posting)
        
        # Generate output file
        output_filename = f'enhanced_resume.{output_format}'
        output_path = os.path.join(self.temp_dir, output_filename)
        
        if output_format == 'pdf':
            self.create_pdf(enhanced_text, output_path)
        elif output_format in ['doc', 'docx']:
            self.create_docx(enhanced_text, output_path)
        else:
            raise Exception(f"Unsupported output format: {output_format}")
        
        return output_path

